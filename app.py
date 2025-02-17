# # Import necessary libraries
# import streamlit as st
# import pdfplumber
# from docx import Document
# import re
# import os
# from openai import OpenAI
# from dotenv import load_dotenv  # Load environment variables

# # Load API key from .env file
# load_dotenv()
# API_KEY = os.getenv("OPENROUTER_API_KEY")

# # OpenRouter API Configuration
# client = OpenAI(
#     base_url="https://openrouter.ai/api/v1",
#     api_key=API_KEY,
# )

# # Function to extract text from PDF
# def extract_text_from_pdf(file):
#     with pdfplumber.open(file) as pdf:
#         return "".join(page.extract_text() or "" for page in pdf.pages)

# # Function to extract text from DOCX
# def extract_text_from_docx(file):
#     doc = Document(file)
#     return "\n".join(para.text for para in doc.paragraphs)

# # Function to extract candidate name and email
# def extract_info(text):
#     email = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
#     email = email[0] if email else "Not Found"
#     name = text.split('\n')[0].strip()  # Assume first line is the name
#     return name, email

# # Function to get ATS score using AI
# def get_ats_score_with_ai(text, required_skills):
#     prompt = f"""
#     Evaluate the ATS compatibility of the following resume based on the given required skills.

#     Resume Text:
#     {text}

#     Required Skills:
#     {', '.join(required_skills)}

#     Provide:
#     1. A percentage-based ATS Score (0-100%).
#     2. A short explanation of why this score was given.
#     """

#     try:
#         completion = client.chat.completions.create(
#             model="deepseek/deepseek-r1-distill-llama-70b:free",
#             messages=[{"role": "user", "content": prompt}],
#         )
#         return completion.choices[0].message.content
#     except Exception as e:
#         return f"Failed to calculate ATS score. Error: {e}"

# # Function to get AI resume improvement suggestions
# def get_resume_suggestions(text, required_skills):
#     prompt = f"""
#     The candidate's resume text is: {text}
#     The required skills for the job are: {', '.join(required_skills)}
#     Provide suggestions to improve the resume to better match the required skills.
#     """

#     try:
#         completion = client.chat.completions.create(
#             model="deepseek/deepseek-r1-distill-llama-70b:free",
#             messages=[{"role": "user", "content": prompt}],
#         )
#         return completion.choices[0].message.content
#     except Exception as e:
#         return f"Failed to get suggestions. Error: {e}"

# # Streamlit App
# def main():
#     st.title("Resume Parser and ATS Score Analyzer")

#     # Upload Resume
#     uploaded_file = st.file_uploader("Upload Resume (PDF or DOCX)", type=["pdf", "docx"])

#     if uploaded_file:
#         # Extract text from the uploaded file
#         text = extract_text_from_pdf(uploaded_file) if uploaded_file.name.endswith('.pdf') else extract_text_from_docx(uploaded_file)

#         # Extract candidate info
#         name, email = extract_info(text)

#         # Display extracted info
#         st.write("### Resume Summary")
#         st.write(f"**Name:** {name}")
#         st.write(f"**Email:** {email}")

#         # Input Required Skills
#         required_skills = st.text_input("Enter Required Skills (comma separated):")
#         required_skills = [skill.strip() for skill in required_skills.split(",")] if required_skills else []

#         if required_skills:
#             # Get AI-calculated ATS score
#             st.write("### AI-Powered ATS Score Analysis")
#             ats_analysis = get_ats_score_with_ai(text, required_skills)
#             st.write(ats_analysis)

#             # Get AI resume improvement suggestions
#             st.write("### Resume Improvement Suggestions")
#             suggestions = get_resume_suggestions(text, required_skills)
#             st.write(suggestions)

#     # Reset Button
#     if st.button("Reset"):
#         st.experimental_rerun()  # Restart the script

# # Run the app
# if __name__ == "__main__":
#     main()
 

















# import streamlit as st
# import pdfplumber
# from docx import Document
# import re
# import os
# from openai import OpenAI
# from dotenv import load_dotenv  
# from PyPDF2 import PdfReader

# # Load API key from .env file
# load_dotenv()
# API_KEY = os.getenv("OPENROUTER_API_KEY")

# # OpenRouter API Configuration
# client = OpenAI(
#     base_url="https://openrouter.ai/api/v1",
#     api_key=API_KEY,
# )

# # Function to extract text from PDF
# def extract_text_from_pdf(file):
#     with pdfplumber.open(file) as pdf:
#         return "".join(page.extract_text() or "" for page in pdf.pages)

# # Function to extract text from DOCX
# def extract_text_from_docx(file):
#     doc = Document(file)
#     return "\n".join(para.text for para in doc.paragraphs)

# # Function to extract candidate name and email
# def extract_info(text):
#     email = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
#     email = email[0] if email else "Not Found"
#     name = text.split('\n')[0].strip()  # Assume first line is the name
#     return name, email

# # Function to get ATS score using AI
# def get_ats_score_with_ai(text, required_skills):
#     prompt = f"""
#     Evaluate the ATS compatibility of the following resume based on the given required skills.

#     Resume Text:
#     {text}

#     Required Skills:
#     {', '.join(required_skills)}

#     Provide:
#     1. A percentage-based ATS Score (0-100%).
#     2. A short explanation of why this score was given.
#     """
    
#     try:
#         completion = client.chat.completions.create(
#             model="deepseek/deepseek-r1-distill-llama-70b:free",
#             messages=[{"role": "user", "content": prompt}],
#         )
#         return completion.choices[0].message.content
#     except Exception as e:
#         return f"Failed to calculate ATS score. Error: {e}"

# # Function to get AI resume improvement suggestions
# def get_resume_suggestions(text, required_skills):
#     prompt = f"""
#     The candidate's resume text is: {text}
#     The required skills for the job are: {', '.join(required_skills)}
#     Provide suggestions to improve the resume to better match the required skills.
#     """
    
#     try:
#         completion = client.chat.completions.create(
#             model="deepseek/deepseek-r1-distill-llama-70b:free",
#             messages=[{"role": "user", "content": prompt}],
#         )
#         return completion.choices[0].message.content
#     except Exception as e:
#         return f"Failed to get suggestions. Error: {e}"

# # Function to fetch job recommendations based on resume text
# def get_job_recommendations(text):
#     prompt = f"""
#     Based on the following resume text, suggest top 3 job roles that fit the candidate's experience and skills:
#     {text}
#     """
    
#     try:
#         completion = client.chat.completions.create(
#             model="deepseek/deepseek-r1-distill-llama-70b:free",
#             messages=[{"role": "user", "content": prompt}],
#         )
#         return completion.choices[0].message.content
#     except Exception as e:
#         return f"Failed to fetch job recommendations. Error: {e}"

# # Streamlit App
# def main():
#     st.set_page_config(layout="wide")  # Wide layout for better visualization
#     st.title("📄 Resume Parser and ATS Score Analyzer")

#     # File Upload
#     uploaded_file = st.file_uploader("Upload Resume (PDF or DOCX)", type=["pdf", "docx"])

#     if uploaded_file:
#         col1, col2 = st.columns([1, 2])  # Left: PDF Preview, Right: Analysis

#         with col1:
#             st.subheader("📑 Resume Preview")
#             if uploaded_file.name.endswith('.pdf'):
#                 pdf_reader = PdfReader(uploaded_file)
#                 for page in pdf_reader.pages:
#                     st.text(page.extract_text())
#             else:
#                 st.write("(Preview not available for DOCX files)")
        
#         with col2:
#             # Extract text
#             text = extract_text_from_pdf(uploaded_file) if uploaded_file.name.endswith('.pdf') else extract_text_from_docx(uploaded_file)
#             name, email = extract_info(text)
            
#             st.subheader("🔍 Extracted Information")
#             st.write(f"**Name:** {name}")
#             st.write(f"**Email:** {email}")
            
#             # Required Skills Input
#             required_skills = st.text_input("🎯 Enter Required Skills (comma separated):")
#             required_skills = [skill.strip() for skill in required_skills.split(",")] if required_skills else []
            
#             if required_skills:
#                 st.subheader("📊 AI-Powered ATS Score Analysis")
#                 ats_analysis = get_ats_score_with_ai(text, required_skills)
#                 st.write(ats_analysis)
                
#                 st.subheader("🛠 Resume Improvement Suggestions")
#                 suggestions = get_resume_suggestions(text, required_skills)
#                 st.write(suggestions)
                
#                 st.subheader("💼 Job Recommendations")
#                 job_recommendations = get_job_recommendations(text)
#                 st.write(job_recommendations)
    
#     # Reset Button
#     if st.button("🔄 Reset"):
#         st.rerun()

# # Run the app
# if __name__ == "__main__":
#     main()




import streamlit as st
import fitz  # Faster PDF processing (PyMuPDF)
from docx import Document
import re
import os
from openai import OpenAI
from dotenv import load_dotenv  
from PyPDF2 import PdfReader
from concurrent.futures import ThreadPoolExecutor  # Multi-threading for speed

# Load API Key Once
load_dotenv()
API_KEY = os.getenv("OPENROUTER_API_KEY")

# Singleton OpenAI client
client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=API_KEY)

# Function to extract text from PDF (Fast)
@st.cache_data
def extract_text_from_pdf(file):
    pdf = fitz.open(stream=file.read(), filetype="pdf")
    return " ".join(page.get_text("text") for page in pdf)

# Function to extract text from DOCX
@st.cache_data
def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join(para.text for para in doc.paragraphs)

# Extract candidate name and email
@st.cache_data
def extract_info(text):
    email = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
    return text.split('\n')[0].strip(), email[0] if email else "Not Found"

# Generic AI API Call (Shared for ATS Score & Suggestions)
@st.cache_data
def call_openai(prompt):
    try:
        response = client.chat.completions.create(
            model="deepseek/deepseek-r1-distill-llama-70b:free",
            messages=[{"role": "user", "content": prompt}],
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Failed to process request. Error: {e}"

# AI ATS Score Analysis
def get_ats_score_with_ai(text, required_skills):
    prompt = f"""
    Evaluate the ATS compatibility of the following resume based on the given required skills.
    Resume Text: {text}
    Required Skills: {', '.join(required_skills)}
    Provide:
    1. A percentage-based ATS Score (0-100%).
    2. A short explanation of why this score was given.
    """
    return call_openai(prompt)

# AI Resume Suggestions
def get_resume_suggestions(text, required_skills):
    prompt = f"""
    The candidate's resume text is: {text}
    The required skills for the job are: {', '.join(required_skills)}
    Provide suggestions to improve the resume to better match the required skills.
    """
    return call_openai(prompt)

# AI Job Recommendations
def get_job_recommendations(text):
    prompt = f"""
    Based on the following resume text, suggest top 3 job roles that fit the candidate's experience and skills:
    {text}
    """
    return call_openai(prompt)

# Streamlit App
def main():
    st.set_page_config(layout="wide")
    st.title("📄 Resume Parser & ATS Analyzer (Fast Mode)")

    uploaded_file = st.file_uploader("Upload Resume (PDF or DOCX)", type=["pdf", "docx"])
    
    if uploaded_file:
        col1, col2 = st.columns([1, 2])

        with col1:
            st.subheader("📑 Resume Preview")
            text = extract_text_from_pdf(uploaded_file) if uploaded_file.name.endswith('.pdf') else extract_text_from_docx(uploaded_file)
            st.text(text[:500])  # Show only first 500 characters

        with col2:
            name, email = extract_info(text)
            st.subheader("🔍 Extracted Information")
            st.write(f"**Name:** {name}")
            st.write(f"**Email:** {email}")

            # Input Required Skills
            required_skills = st.text_input("🎯 Enter Required Skills (comma separated):")
            required_skills = [skill.strip() for skill in required_skills.split(",")] if required_skills else []
            
            if required_skills:
                st.subheader("📊 AI ATS Score & Suggestions")
                
                # Parallel Processing: Speed Up AI Calls
                with ThreadPoolExecutor() as executor:
                    ats_future = executor.submit(get_ats_score_with_ai, text, required_skills)
                    suggestions_future = executor.submit(get_resume_suggestions, text, required_skills)
                    job_future = executor.submit(get_job_recommendations, text)

                    ats_analysis = ats_future.result()
                    suggestions = suggestions_future.result()
                    job_recommendations = job_future.result()

                st.write(ats_analysis)
                st.subheader("🛠 Resume Improvement Suggestions")
                st.write(suggestions)

                st.subheader("💼 Job Recommendations")
                st.write(job_recommendations)

    if st.button("🔄 Reset"):
        st.rerun()

# Run the app
if __name__ == "__main__":
    main()
